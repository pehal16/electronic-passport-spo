from __future__ import annotations

import io

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.schemas.attestation import AttestationSheetRead
from app.services.attestation import ATTENDANCE_RESULT_TITLE

HEADING_BY_TEMPLATE = {
    "diff_credit": "ВЕДОМОСТЬ ДИФФЕРЕНЦИРОВАННОГО ЗАЧЕТА",
    "credit_sheet": "ЗАЧЕТНАЯ ВЕДОМОСТЬ",
    "complex_diff_credit": "ВЕДОМОСТЬ КОМПЛЕКСНОГО ДИФФЕРЕНЦИРОВАННОГО ЗАЧЕТА",
    "complex_exam": "ВЕДОМОСТЬ КОМПЛЕКСНОГО ЭКЗАМЕНА",
}


def build_pdf_bytes(sheet: AttestationSheetRead) -> bytes:
    is_complex_exam = sheet.sheet_template.code == "complex_exam"
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    styles = getSampleStyleSheet()
    heading = HEADING_BY_TEMPLATE.get(sheet.sheet_template.code, sheet.title.upper())
    content = [
        Paragraph(f"<b>{sheet.college_name}</b>", styles["Title"]),
        Spacer(1, 6),
        Paragraph(f"<b>{heading}</b>", styles["Heading2"]),
        Spacer(1, 6),
        Paragraph(f"Группа: {sheet.group_name}", styles["Normal"]),
        Paragraph(f"{sheet.header_label}: {sheet.header_value}", styles["Normal"]),
        Paragraph(f"Дисциплина: {sheet.discipline_display_text}", styles["Normal"]),
        Paragraph(f"Преподаватель: {sheet.teacher_name}", styles["Normal"]),
        Paragraph(f"Дата: {sheet.date}", styles["Normal"]),
        Spacer(1, 8),
    ]

    table_header = ["№", "ФИО студента"]
    if sheet.sheet_template.has_ticket_number:
        table_header.append("Билет/вариант")
    table_header.append("Оценка (цифрой и прописью)" if is_complex_exam else "Оценка")
    if not is_complex_exam:
        table_header.append("Статус")
    table_header.append("Подпись экзаменатора" if is_complex_exam else "Подпись")

    rows = [table_header]
    for row in sheet.rows:
        values = [str(row.row_number), row.student_name_snapshot]
        if sheet.sheet_template.has_ticket_number:
            values.append(row.ticket_number or "")
        values.append(row.grade_text or row.grade_numeric or "")
        if not is_complex_exam:
            values.append(ATTENDANCE_RESULT_TITLE.get(row.attendance_result, row.attendance_result))
        values.append("")
        rows.append(values)

    table = Table(rows, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2f2f2")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fafafa")]),
            ]
        )
    )
    content.append(table)
    content.append(Spacer(1, 12))

    totals = sheet.totals
    top_line = f"Допущено: {totals.admitted}" if is_complex_exam else f"Итого: {totals.total_rows} ({totals.total_rows_words})"
    totals_text = (
        f"Отлично: {totals.excellent}; Хорошо: {totals.good}; "
        f"Удовлетворительно: {totals.satisfactory}; Неудовлетворительно: {totals.unsatisfactory}; "
        f"Не сдавали: {totals.not_submitted}; Не явились: {totals.not_appeared}."
    )
    content.append(Paragraph(top_line, styles["Normal"]))
    content.append(Paragraph(totals_text, styles["Normal"]))
    content.append(Spacer(1, 10))
    content.append(Paragraph("Преподаватель: ____________________", styles["Normal"]))
    content.append(Paragraph("Заведующий отделением: ____________________", styles["Normal"]))

    document.build(content)
    return buffer.getvalue()


def build_docx_bytes(sheet: AttestationSheetRead) -> bytes:
    is_complex_exam = sheet.sheet_template.code == "complex_exam"
    heading = HEADING_BY_TEMPLATE.get(sheet.sheet_template.code, sheet.title.upper())
    doc = Document()
    doc.add_paragraph(sheet.college_name).bold = True
    doc.add_paragraph(heading).bold = True
    doc.add_paragraph(f"Группа: {sheet.group_name}")
    doc.add_paragraph(f"{sheet.header_label}: {sheet.header_value}")
    doc.add_paragraph(f"Дисциплина: {sheet.discipline_display_text}")
    doc.add_paragraph(f"Преподаватель: {sheet.teacher_name}")
    doc.add_paragraph(f"Дата: {sheet.date}")

    header_cells = ["№", "ФИО студента"]
    if sheet.sheet_template.has_ticket_number:
        header_cells.append("Билет/вариант")
    header_cells.append("Оценка (цифрой и прописью)" if is_complex_exam else "Оценка")
    if not is_complex_exam:
        header_cells.append("Статус")
    header_cells.append("Подпись экзаменатора" if is_complex_exam else "Подпись")

    table = doc.add_table(rows=1, cols=len(header_cells))
    table.style = "Table Grid"
    for idx, header in enumerate(header_cells):
        table.cell(0, idx).text = header

    for row in sheet.rows:
        values = [str(row.row_number), row.student_name_snapshot]
        if sheet.sheet_template.has_ticket_number:
            values.append(row.ticket_number or "")
        values.append(row.grade_text or row.grade_numeric or "")
        if not is_complex_exam:
            values.append(ATTENDANCE_RESULT_TITLE.get(row.attendance_result, row.attendance_result))
        values.append("")
        row_cells = table.add_row().cells
        for idx, value in enumerate(values):
            row_cells[idx].text = value

    totals = sheet.totals
    top_line = f"Допущено: {totals.admitted}" if is_complex_exam else f"Итого: {totals.total_rows} ({totals.total_rows_words})"
    totals_text = (
        f"Отлично: {totals.excellent}; Хорошо: {totals.good}; "
        f"Удовлетворительно: {totals.satisfactory}; Неудовлетворительно: {totals.unsatisfactory}; "
        f"Не сдавали: {totals.not_submitted}; Не явились: {totals.not_appeared}."
    )
    doc.add_paragraph(top_line)
    doc.add_paragraph(totals_text)
    doc.add_paragraph("Преподаватель: ____________________")
    doc.add_paragraph("Заведующий отделением: ____________________")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
